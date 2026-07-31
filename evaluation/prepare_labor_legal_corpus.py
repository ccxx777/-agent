"""把劳动合同 A 级 Word 法律资料准备为可核验的 Markdown 与 metadata。

该脚本只处理本地 ``data/legal/labor_contract/raw/a_level`` 中的 DOCX 文件。
它不会推断法律效力，也不会把原文直接写入 Qdrant；未知的来源、制定机关、
施行日期和效力状态统一使用显式占位字符串，等待人工核验后再导入。

示例：

    uv run --with python-docx python evaluation/prepare_labor_legal_corpus.py

生成物（均位于 Git 忽略的 data/）包括：

* normalized/a_level/*.md：保留章节和条文段落的 Markdown；
* metadata/a_level_documents.jsonl：每份资料一行的机器可读 metadata；
* manifests/legal_labor_a_v1_preparation.json：本次准备批次的汇总清单；
* metadata/a_level_manual_review.md：人工补齐字段的操作说明。
"""

from __future__ import annotations

import argparse
import hashlib
import json
import re
from collections.abc import Iterable
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

FORMAT_VERSION = 1
PLACEHOLDER = "PENDING_MANUAL_VERIFICATION"
OFFICIAL_URL_PLACEHOLDER = "UNVERIFIED_OFFICIAL_URL"

# 这些字段不由本地 DOCX 转换推断。重新执行 --overwrite 时必须保留人工或
# 官方数据库回填的值，避免已核验 metadata 被新的占位符覆盖。
EXTERNAL_METADATA_FIELDS = (
    "document_type",
    "issuing_authority",
    "jurisdiction",
    "national_applicability",
    "publication_date",
    "effective_date",
    "effective_date_source",
    "effective_date_source_url",
    "effective_date_verified_at",
    "amendment_or_repeal_status",
    "official_url",
    "license_status",
    "pii_status",
    "review_status",
    "official_source_id",
    "official_metadata_source",
    "official_metadata_checked_at",
    "official_status_code",
    "content_match_status",
)

_FILENAME_DATE = re.compile(r"_(\d{8})$")
_CHAPTER = re.compile(r"^第[一二三四五六七八九十百零〇0-9]+章")
_SECTION = re.compile(r"^第[一二三四五六七八九十百零〇0-9]+节")
_ARTICLE = re.compile(r"^第[一二三四五六七八九十百零〇0-9]+条")

# 当前最小包的稳定英文 ID。未知文件不应被默默塞进 Collection，因此会使用
# 哈希派生的 ID，并仍然保留待人工核验状态。
KNOWN_DOCUMENT_IDS = {
    "中华人民共和国劳动合同法": "labor_contract_law_20121228",
    "中华人民共和国劳动合同法实施条例": "labor_contract_law_implementation_regulation_20080918",
    "中华人民共和国劳动法": "labor_law_20181229",
    "中华人民共和国社会保险法": "social_insurance_law_20181229",
    "中华人民共和国劳动争议调解仲裁法": "labor_dispute_mediation_arbitration_law_20071229",
    "最高人民法院关于审理劳动争议案件适用法律问题的解释（一）": "spc_labor_dispute_interpretation_1_20201229",
    "最高人民法院关于审理劳动争议案件适用法律问题的解释（二）": "spc_labor_dispute_interpretation_2_20250731",
}


@dataclass(frozen=True)
class PreparedDocument:
    """一份准备完毕但尚未人工核验的法律资料。"""

    doc_id: str
    title: str
    raw_path: Path
    markdown_path: Path
    raw_sha256: str
    markdown_sha256: str
    source_filename_date: str
    paragraph_count: int
    table_count: int


def _sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def _clean_text(text: str) -> str:
    """保留段落边界，仅清理 DOCX 中常见的不可见空白。"""

    text = text.replace("\ufeff", "").replace("\u200b", "")
    text = text.replace("\xa0", " ").replace("\u3000", " ")
    return re.sub(r"[ \t]+", " ", text).strip()


def _title_and_filename_date(path: Path) -> tuple[str, str]:
    match = _FILENAME_DATE.search(path.stem)
    if match:
        raw_date = match.group(1)
        source_filename_date = f"{raw_date[:4]}-{raw_date[4:6]}-{raw_date[6:]}"
        return path.stem[: match.start()], source_filename_date
    return path.stem, "UNVERIFIED_FILENAME_DATE"


def _iter_blocks(document: Document) -> Iterable[Paragraph | Table]:
    """按原始文档顺序输出段落和表格，避免表格被静默丢弃。"""

    for child in document.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, document)
        elif isinstance(child, CT_Tbl):
            yield Table(child, document)


def _table_markdown(table: Table) -> list[str]:
    rows = [
        [_clean_text(cell.text).replace("|", "\\|").replace("\n", "<br>") for cell in row.cells]
        for row in table.rows
    ]
    rows = [row for row in rows if any(cell for cell in row)]
    if not rows:
        return []
    width = max(len(row) for row in rows)
    padded = [row + [""] * (width - len(row)) for row in rows]
    header = padded[0]
    separator = ["---"] * width
    lines = ["| " + " | ".join(header) + " |", "| " + " | ".join(separator) + " |"]
    lines.extend("| " + " | ".join(row) + " |" for row in padded[1:])
    return lines


def _paragraph_markdown(text: str) -> str:
    if _CHAPTER.match(text):
        return f"## {text}"
    if _SECTION.match(text):
        return f"### {text}"
    if _ARTICLE.match(text):
        return f"#### {text}"
    return text


def _doc_id(title: str, raw_sha256: str) -> str:
    return KNOWN_DOCUMENT_IDS.get(title, f"unverified_legal_document_{raw_sha256[:12]}")


def _relative(path: Path, base: Path) -> str:
    return path.resolve().relative_to(base.resolve()).as_posix()


def _metadata(record: PreparedDocument, *, base: Path) -> dict[str, Any]:
    """不把文件名日期误写为生效日期；所有法律属性均等待人工确认。"""

    return {
        "format_version": FORMAT_VERSION,
        "doc_id": record.doc_id,
        "title": record.title,
        "source_level": "A",
        "target_collection": "legal_labor_a_v1",
        "document_type": PLACEHOLDER,
        "issuing_authority": PLACEHOLDER,
        "jurisdiction": PLACEHOLDER,
        "national_applicability": PLACEHOLDER,
        "publication_date": PLACEHOLDER,
        "effective_date": PLACEHOLDER,
        "amendment_or_repeal_status": PLACEHOLDER,
        "official_url": OFFICIAL_URL_PLACEHOLDER,
        "license_status": PLACEHOLDER,
        "pii_status": "NOT_APPLICABLE_FOR_STATUTE_TEXT_PENDING_REVIEW",
        "review_status": "PENDING_MANUAL_REVIEW",
        "source_filename_date": record.source_filename_date,
        "raw_file": _relative(record.raw_path, base),
        "normalized_markdown": _relative(record.markdown_path, base),
        "raw_sha256": record.raw_sha256,
        "normalized_markdown_sha256": record.markdown_sha256,
        "paragraph_count": record.paragraph_count,
        "table_count": record.table_count,
        "parser": "python-docx",
    }


def _load_existing_metadata(path: Path) -> dict[str, dict[str, Any]]:
    """按 doc_id 读取已有 JSONL；损坏行直接拒绝，避免静默丢失核验结果。"""

    if not path.is_file():
        return {}
    records: dict[str, dict[str, Any]] = {}
    for line_number, line in enumerate(path.read_text(encoding="utf-8").splitlines(), start=1):
        if not line.strip():
            continue
        try:
            value = json.loads(line)
        except json.JSONDecodeError as error:
            raise ValueError(f"已有 metadata 第 {line_number} 行不是有效 JSON") from error
        doc_id = value.get("doc_id")
        if not isinstance(doc_id, str) or not doc_id:
            raise ValueError(f"已有 metadata 第 {line_number} 行缺少 doc_id")
        if doc_id in records:
            raise ValueError(f"已有 metadata 存在重复 doc_id：{doc_id}")
        records[doc_id] = value
    return records


def _load_existing_manifest(path: Path) -> dict[str, Any]:
    """重建文本产物时保留已经记录的官方 metadata 核验批次。"""

    if not path.is_file():
        return {}
    try:
        value = json.loads(path.read_text(encoding="utf-8"))
    except json.JSONDecodeError as error:
        raise ValueError("已有 manifest 不是有效 JSON") from error
    if not isinstance(value, dict):
        raise TypeError("已有 manifest 必须是 JSON 对象")
    return value


def _merge_external_metadata(
    generated: dict[str, Any], existing: dict[str, Any] | None
) -> dict[str, Any]:
    """只保留不应从文件名或 DOCX 内容推断的外部核验字段。"""

    if not existing:
        return generated
    result = dict(generated)
    for field in EXTERNAL_METADATA_FIELDS:
        if field in existing:
            result[field] = existing[field]
    return result


def _front_matter(metadata: dict[str, Any]) -> str:
    """以 JSON front matter 写入 Markdown，避免 YAML 引号和中文转义歧义。"""

    selected = {
        key: metadata[key]
        for key in (
            "format_version",
            "doc_id",
            "title",
            "source_level",
            "target_collection",
            "document_type",
            "issuing_authority",
            "jurisdiction",
            "national_applicability",
            "publication_date",
            "effective_date",
            "effective_date_source",
            "effective_date_source_url",
            "effective_date_verified_at",
            "amendment_or_repeal_status",
            "official_url",
            "review_status",
            "source_filename_date",
            "raw_file",
            "raw_sha256",
        )
        if key in metadata
    }
    return "---json\n" + json.dumps(selected, ensure_ascii=False, indent=2) + "\n---"


def _render_markdown(title: str, blocks: list[Paragraph | Table], metadata: dict[str, Any]) -> tuple[str, int, int]:
    lines = [_front_matter(metadata), "", f"# {title}", ""]
    lines.extend(
        [
            "> 本文件由本地 DOCX 自动转换。法律效力、制定机关、官方链接和施行日期尚未人工核验；",
            "> 不应据此直接作出正式法律结论或导入生产资料库。",
            "",
            "## 原文正文",
            "",
        ]
    )
    paragraph_count = 0
    table_count = 0
    title_seen = False
    for block in blocks:
        if isinstance(block, Paragraph):
            text = _clean_text(block.text)
            if not text:
                continue
            # 正文第一行通常重复文档标题，Markdown 已有 H1，因此只保留一次。
            if text == title and not title_seen:
                title_seen = True
                continue
            title_seen = title_seen or text == title
            lines.extend([_paragraph_markdown(text), ""])
            paragraph_count += 1
            continue
        table_lines = _table_markdown(block)
        if table_lines:
            lines.extend(table_lines)
            lines.append("")
        table_count += 1
    return "\n".join(lines).rstrip() + "\n", paragraph_count, table_count


def _manual_review_markdown(records: list[PreparedDocument]) -> str:
    rows = [
        "# A 级法律资料人工核验清单",
        "",
        "以下文件的 SHA-256、文件名日期和本地路径已由脚本生成。",
        "部分字段可能已由官方数据库回填；所有仍标为待核验的字段，以及正文一致性和适用范围，",
        "都必须在进入 `legal_labor_a_v1` 之前完成人工复核。",
        "",
        "| 文档 ID | 文件名 | 需要人工补齐 |",
        "|---|---|---|",
    ]
    for record in records:
        rows.append(
            "| "
            f"`{record.doc_id}` | `{record.raw_path.name}` | "
            "官方 URL、制定机关、文种、公布/施行日期、效力状态、全国适用性、来源授权、复核人 |"
        )
    rows.extend(
        [
            "",
            "## 必须确认的边界",
            "",
            "- 文件名中的日期只作为 `source_filename_date` 保存，不能自动当作 `effective_date`。",
            "- 同名法律的版本必须以官方来源、文本内容和 SHA-256 联合确认。",
            "- 司法解释、修订法规和已废止条款必须写明衔接关系；未知时保持待核验。",
            "- 法律原文仅进入独立的 A 级资料库，不能写入 `rag_chunks`、watsonxDocsQA Collection 或用户合同库。",
        ]
    )
    return "\n".join(rows) + "\n"


def prepare(*, raw_dir: Path, base_dir: Path, overwrite: bool) -> dict[str, Any]:
    if not raw_dir.is_dir():
        raise FileNotFoundError(f"A 级原始资料目录不存在：{raw_dir}")
    source_files = sorted(raw_dir.glob("*.docx"))
    if not source_files:
        raise FileNotFoundError(f"未在 {raw_dir} 找到 DOCX 文件")

    normalized_dir = base_dir / "normalized" / "a_level"
    metadata_dir = base_dir / "metadata"
    manifest_dir = base_dir / "manifests"
    metadata_file = metadata_dir / "a_level_documents.jsonl"
    manifest_file = manifest_dir / "legal_labor_a_v1_preparation.json"
    review_file = metadata_dir / "a_level_manual_review.md"

    outputs = [metadata_file, manifest_file, review_file]
    outputs.extend(normalized_dir / f"{path.stem}.md" for path in source_files)
    if not overwrite:
        existing = [path for path in outputs if path.exists()]
        if existing:
            joined = ", ".join(str(path) for path in existing)
            raise FileExistsError(f"输出已存在；如需重建请添加 --overwrite：{joined}")

    existing_metadata = _load_existing_metadata(metadata_file) if overwrite else {}
    existing_manifest = _load_existing_manifest(manifest_file) if overwrite else {}

    normalized_dir.mkdir(parents=True, exist_ok=True)
    metadata_dir.mkdir(parents=True, exist_ok=True)
    manifest_dir.mkdir(parents=True, exist_ok=True)

    records: list[PreparedDocument] = []
    metadata_rows: list[dict[str, Any]] = []
    for raw_path in source_files:
        title, source_filename_date = _title_and_filename_date(raw_path)
        raw_sha256 = _sha256(raw_path)
        doc_id = _doc_id(title, raw_sha256)
        markdown_path = normalized_dir / f"{doc_id}.md"
        document = Document(raw_path)
        provisional = PreparedDocument(
            doc_id=doc_id,
            title=title,
            raw_path=raw_path,
            markdown_path=markdown_path,
            raw_sha256=raw_sha256,
            markdown_sha256="PENDING_MARKDOWN_RENDER",
            source_filename_date=source_filename_date,
            paragraph_count=0,
            table_count=0,
        )
        provisional_metadata = _merge_external_metadata(
            _metadata(provisional, base=base_dir),
            existing_metadata.get(doc_id),
        )
        markdown, paragraph_count, table_count = _render_markdown(
            title,
            list(_iter_blocks(document)),
            provisional_metadata,
        )
        markdown_path.write_text(markdown, encoding="utf-8")
        final_record = PreparedDocument(
            doc_id=doc_id,
            title=title,
            raw_path=raw_path,
            markdown_path=markdown_path,
            raw_sha256=raw_sha256,
            markdown_sha256=_sha256(markdown_path),
            source_filename_date=source_filename_date,
            paragraph_count=paragraph_count,
            table_count=table_count,
        )
        records.append(final_record)
        metadata_rows.append(
            _merge_external_metadata(
                _metadata(final_record, base=base_dir),
                existing_metadata.get(doc_id),
            )
        )

    metadata_file.write_text(
        "".join(json.dumps(row, ensure_ascii=False, sort_keys=True) + "\n" for row in metadata_rows),
        encoding="utf-8",
    )
    review_file.write_text(_manual_review_markdown(records), encoding="utf-8")
    manifest = {
        "format_version": FORMAT_VERSION,
        "corpus": "labor_contract_a_level",
        "target_collection": "legal_labor_a_v1",
        "prepared_at": datetime.now(timezone.utc).isoformat(),
        "source_directory": _relative(raw_dir, base_dir),
        "document_count": len(records),
        "documents": metadata_rows,
        "status": "PREPARED_PENDING_MANUAL_METADATA_REVIEW",
        "blocked_actions": [
            "Do not ingest into legal_labor_a_v1 until every PENDING_MANUAL_VERIFICATION field is reviewed.",
            "Do not use these documents as formal legal authority before source and effective-date validation.",
        ],
    }
    if "official_metadata_verification" in existing_manifest:
        manifest["official_metadata_verification"] = existing_manifest[
            "official_metadata_verification"
        ]
        manifest["status"] = existing_manifest.get(
            "status",
            "OFFICIAL_METADATA_FETCHED_PENDING_CONTENT_AND_SCOPE_REVIEW",
        )
    if "manual_official_overrides" in existing_manifest:
        # 补充的官方正文核验记录与 document metadata 同样属于外部事实；
        # --overwrite 只重建 DOCX 派生产物，不能丢弃其来源和核验时间。
        manifest["manual_official_overrides"] = existing_manifest[
            "manual_official_overrides"
        ]
    manifest_file.write_text(json.dumps(manifest, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    return {
        "status": "prepared",
        "documents": len(records),
        "normalized_directory": str(normalized_dir),
        "metadata": str(metadata_file),
        "manual_review": str(review_file),
        "manifest": str(manifest_file),
    }


def parse_args() -> argparse.Namespace:
    workspace_root = Path(__file__).resolve().parents[1]
    default_base = workspace_root / "data" / "legal" / "labor_contract"
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--base", type=Path, default=default_base, help="劳动合同资料根目录")
    parser.add_argument(
        "--raw",
        type=Path,
        default=None,
        help="A 级 DOCX 目录；默认使用 <base>/raw/a_level",
    )
    parser.add_argument("--overwrite", action="store_true", help="允许覆盖已有准备产物")
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    base_dir = args.base.resolve()
    raw_dir = (args.raw or base_dir / "raw" / "a_level").resolve()
    try:
        result = prepare(raw_dir=raw_dir, base_dir=base_dir, overwrite=args.overwrite)
    except (FileNotFoundError, FileExistsError, TypeError, ValueError) as error:
        print(f"[FAIL] {error}")
        return 1
    print(json.dumps(result, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
